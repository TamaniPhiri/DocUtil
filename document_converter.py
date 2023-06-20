import streamlit as st
from PyPDF2 import PdfWriter, PdfReader
from docx import Document
import os

def convert_to_pdf(docx_file):
    pdf_file = f"{os.path.splitext(docx_file)[0]}.pdf"

    document = Document(docx_file)
    paragraphs = document.paragraphs

    pdf_writer = PdfWriter()
    for paragraph in paragraphs:
        pdf_writer.add_page(paragraph._element)

    with open(pdf_file, "wb") as f:
        pdf_writer.write(f)

    return pdf_file

def convert_to_docx(pdf_file):
    docx_file = f"{os.path.splitext(pdf_file)[0]}.docx"

    pdf_reader = PdfReader(pdf_file)
    pdf_pages = [pdf_reader.pages[page_num].extract_text() for page_num in range(len(pdf_reader.pages))]

    doc = Document()
    for page in pdf_pages:
        doc.add_paragraph(page)

    doc.save(docx_file)
    return docx_file

def main():
    st.title("Document Converter")

    uploaded_file = st.file_uploader("Upload a document", type=["docx", "pdf"])
    convert_format = st.selectbox("Select the conversion format", ["DOCX to PDF", "PDF to DOCX"])

    if uploaded_file is not None:
        file_contents = uploaded_file.getvalue()
        file_name = uploaded_file.name

        with open(file_name, "wb") as f:
            f.write(file_contents)

        status_placeholder = st.empty()
        status_placeholder.write("Converting...")

        if convert_format == "DOCX to PDF" and file_name.endswith(".docx"):
            converted_file = convert_to_pdf(file_name)
        elif convert_format == "PDF to DOCX" and file_name.endswith(".pdf"):
            converted_file = convert_to_docx(file_name)
        else:
            st.error("Invalid file format or conversion option.")
            return

        status_placeholder.empty()
        st.success("Conversion complete!")

        # Provide the download button for the converted file
        st.markdown(f"**Download converted file:**")
        download_button_str = f"Download {os.path.basename(converted_file)}"
        st.download_button(download_button_str, converted_file)

if __name__ == "__main__":
    main()